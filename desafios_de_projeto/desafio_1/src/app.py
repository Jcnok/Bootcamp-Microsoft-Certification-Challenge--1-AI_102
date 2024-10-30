# main.py
import streamlit as st
from streamlit_lottie import st_lottie
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
import os
import logging
from datetime import datetime
from docx import Document
from io import BytesIO
import uuid
import json

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Carregar variáveis de ambiente
load_dotenv()

def carregar_lottie(url: str):
    """Carrega animação do Lottie Files."""
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

class TradutorArtigos:
    """Classe para tradução de artigos usando Azure OpenAI."""
    
    def __init__(self):
        self.api_key = os.getenv("AZURE_OPENAI_KEY")
        self.endpoint = "https://openai-dio-boot-east1.openai.azure.com/openai/deployments/gpt-4o-mini/chat/completions?api-version=2024-02-15-preview"

    def extrair_texto(self, url: str) -> str:
        try:
            response = requests.get(url)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
                
            return soup.get_text(" ", strip=True)
            
        except requests.RequestException as e:
            logger.error(f"Falha ao buscar URL: {e}")
            raise

    def traduzir_texto(self, texto: str, idioma_destino: str) -> str:
        headers = {
            "Content-Type": "application/json",
            "api-key": self.api_key,
        }
        
        payload = {
            "messages": [
                {
                    "role": "system",
                    "content": [{"type": "text", "text": "Você atua como tradutor de textos"}]
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": f"traduza: {texto} para o idioma {idioma_destino} e responda apenas com a tradução no formato markdown"
                        }
                    ]
                },    
            ],
            "temperature": 0.7,
            "top_p": 0.95,
            "max_tokens": 900
        }
        
        try:
            response = requests.post(self.endpoint, headers=headers, json=payload)
            response.raise_for_status()
            return response.json()['choices'][0]['message']['content']
            
        except requests.RequestException as e:
            logger.error(f"Falha na requisição da API: {e}")
            raise

class TradutorDocumentos:
    """Classe para tradução de documentos Word usando Azure Translator."""
    
    def __init__(self):
        self.api_key = os.getenv("TRANSLATOR_API_KEY")
        self.endpoint = os.getenv("TRANSLATOR_ENDPOINT")
        self.location = os.getenv("TRANSLATOR_LOCATION")

    def traduzir_texto(self, texto, idioma_origem, idioma_destino):
        path = '/translate'
        constructed_url = self.endpoint + path

        params = {
            'api-version': '3.0',
            'from': idioma_origem,
            'to': [idioma_destino]
        }

        headers = {
            'Ocp-Apim-Subscription-Key': self.api_key,
            'Ocp-Apim-Subscription-Region': self.location,
            'Content-type': 'application/json',
            'X-ClientTraceId': str(uuid.uuid4())
        }

        body = [{'text': texto}]
        response = requests.post(constructed_url, params=params, headers=headers, json=body)
        
        if response.status_code == 200:
            return response.json()[0]['translations'][0]['text']
        else:
            raise Exception(f"Erro na tradução: {response.status_code}")

def pagina_inicial():
    """Página inicial do aplicativo."""
    st.title("🌍 Tradutor Multifuncional")
    
    # Carregar animação
    lottie_translate = carregar_lottie("https://lottie.host/cdc7f167-7c8a-4b67-9a5f-3db936a8cb8d/IB6FNEMBAR.json")
    st_lottie(lottie_translate, height=300)
    
    st.markdown("""
    ## Bem-vindo ao Tradutor Multifuncional!
    
    Nossa plataforma oferece duas poderosas ferramentas de tradução:
    
    ### 📰 Tradutor de Artigos
    Traduza artigos da web mantendo a formatação markdown e estrutura original.
    
    ### 📄 Tradutor de Documentos
    Traduza documentos Word com suporte a múltiplos idiomas e download direto.
    
    ### Tecnologias Utilizadas
    
    #### Tradutor de Artigos:
    - Azure OpenAI GPT-4
    - BeautifulSoup4
    - Requests
    - Streamlit
    
    #### Tradutor de Documentos:
    - Azure Translator API
    - python-docx
    - Streamlit
    - UUID
    
    ### Sobre o Projeto
    Este projeto foi desenvolvido para facilitar a tradução de conteúdo em diferentes formatos,
    oferecendo uma interface intuitiva e resultados profissionais.
    """)

def pagina_tradutor_artigos():
    """Página do tradutor de artigos."""
    st.title("📰 Tradutor de Artigos")
    
    # Carregar animação
    lottie_article = carregar_lottie("https://lottie.host/cdc7f167-7c8a-4b67-9a5f-3db936a8cb8d/IB6FNEMBAR.json")
    st_lottie(lottie_article, height=200)
    
    st.markdown("""
    ### Como funciona?
    1. Cole a URL do artigo que deseja traduzir
    2. Selecione o idioma de destino
    3. Clique em traduzir
    4. Baixe o resultado em formato markdown
    """)
    
    tradutor = TradutorArtigos()
    
    url = st.text_input("🔗 URL do Artigo")
    
    idiomas = {
        "Português": "português",
        "Inglês": "english",
        "Espanhol": "español",
        "Francês": "français",
        "Alemão": "deutsch",
        "Italiano": "italiano"
    }
    
    idioma_destino = st.selectbox("🌍 Idioma de Destino", list(idiomas.keys()))
    
    if st.button("🔄 Traduzir Artigo"):
        try:
            with st.spinner("Processando artigo..."):
                texto = tradutor.extrair_texto(url)
                traducao = tradutor.traduzir_texto(texto, idiomas[idioma_destino])
                
                st.markdown("### Resultado da Tradução")
                st.markdown(traducao)
                
                # Botão de download
                st.download_button(
                    label="📥 Baixar Tradução",
                    data=traducao.encode('utf-8'),
                    file_name=f"traducao_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                    mime="text/markdown"
                )
                
        except Exception as e:
            st.error(f"Ocorreu um erro: {str(e)}")

def pagina_tradutor_documentos():
    """Página do tradutor de documentos."""
    st.title("📄 Tradutor de Documentos Word")
    
    # Carregar animação
    lottie_doc = carregar_lottie("https://lottie.host/cdc7f167-7c8a-4b67-9a5f-3db936a8cb8d/IB6FNEMBAR.json")
    st_lottie(lottie_doc, height=200)
    
    st.markdown("""
    ### Como funciona?
    1. Faça upload do seu documento Word
    2. Selecione os idiomas de origem e destino
    3. Clique em traduzir
    4. Baixe o documento traduzido
    """)
    
    tradutor = TradutorDocumentos()
    
    arquivo = st.file_uploader("📎 Carregar arquivo Word", type=["docx"])
    
    idiomas = {
        'Inglês': 'en',
        'Francês': 'fr',
        'Espanhol': 'es',
        'Alemão': 'de',
        'Português': 'pt',
        'Italiano': 'it'
    }
    
    col1, col2 = st.columns(2)
    with col1:
        idioma_origem = st.selectbox("🔤 Idioma de origem", list(idiomas.keys()))
    with col2:
        idioma_destino = st.selectbox("🔤 Idioma de destino", list(idiomas.keys()))

    if st.button("🔄 Traduzir"):
        if arquivo is not None:
            try:
                with st.spinner("Traduzindo documento..."):
                    # Carregar documento
                    doc = Document(arquivo)
                    texto_original = "\n".join([paragrafo.text for paragrafo in doc.paragraphs])
                    
                    # Traduzir
                    traducao = tradutor.traduzir_texto(
                        texto_original,
                        idiomas[idioma_origem],
                        idiomas[idioma_destino]
                    )
                    
                    # Criar novo documento
                    doc_traduzido = Document()
                    for linha in traducao.split('\n'):
                        doc_traduzido.add_paragraph(linha)
                    
                    # Preparar para download
                    buffer = BytesIO()
                    doc_traduzido.save(buffer)
                    buffer.seek(0)
                    
                    # Botão de download
                    st.download_button(
                        label="📥 Baixar Documento Traduzido",
                        data=buffer,
                        file_name=f"traducao_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
            except Exception as e:
                st.error(f"Ocorreu um erro: {str(e)}")
        else:
            st.error("Por favor, carregue um arquivo Word.")

def main():
    """Função principal do aplicativo."""
    st.set_page_config(page_title="Tradutor Multifuncional", layout="wide")
    
    # Menu lateral
    st.sidebar.title("📚 Menu")
    paginas = {
        "🏠 Página Inicial": pagina_inicial,
        "📰 Tradutor de Artigos": pagina_tradutor_artigos,
        "📄 Tradutor de Documentos": pagina_tradutor_documentos
    }
    
    escolha = st.sidebar.radio("Navegação", list(paginas.keys()))
    
    # Rodapé do menu
    st.sidebar.markdown("---")
    st.sidebar.markdown("""
    ### Sobre
    Desenvolvido com ❤️ usando:
    - Streamlit
    - Azure OpenAI
    - Azure Translator
    - Python
    
    © 2024 Todos os direitos reservados
    """)
    
    # Renderizar página selecionada
    paginas[escolha]()

if __name__ == "__main__":
    main()